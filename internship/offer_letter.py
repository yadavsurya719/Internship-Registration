from docx import Document
from datetime import datetime
from docx2pdf import convert
import mysql.connector
import pymysql

mydb = mysql.connector.connect(host='localhost', user='root', password='surya7993290146', database='data')
mycursor = mydb.cursor()
print("Database is connected")

def offer_letter(full_name, job_role, start_date, salary, Hr_name):
    # Create a new Document
    doc = Document()

    # Add a heading
    doc.add_heading('Offer letter', level=1)

    # Add the current date
    current_date = datetime.now().strftime('%Y-%m-%d')
    doc.add_paragraph(f'Date: {current_date}')

    # Add employee details
    doc.add_paragraph(f'Dear {full_name},')
    doc.add_paragraph(f'We at {full_name} are delighted to inform you that you have been selected to our company as a {job_role}. Your joining date is {start_date} and you have to report to the office at 9:00 AM. Your salary will be {salary} per month.')

    # Calculate the number of days1
    doc.add_paragraph('Regards,')

    # Add a closing statement
    doc.add_paragraph(f'{Hr_name}[HR], Pragmatiq Systems')

    # Add a closing salutation
    doc_filename = f'Offer_Letter_{full_name.replace(" ","_")}.docx'
    doc.save(doc_filename)

    # Save the document
    doc_filename = f'Offer_Letter{full_name}.docx'
    doc.save(doc_filename)

    # Convert the document to PDF
    convert(doc_filename)


# ..
user_id=input("Enter User id:")
# Take user input for user id
mycursor.execute("SELECT * FROM data.internship where id=%s", (user_id,))

results = mycursor.fetchall()

for result in results:
    offer_letter(str(result[0]), result[1], result[2], result[3], result[4])


mydb.close()
