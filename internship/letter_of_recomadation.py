from docx import Document
from datetime import datetime
from docx2pdf import convert
import mysql.connector
import pymysql

mydb = mysql.connector.connect(host='localhost', user='root', password='surya7993290146', database='data')
mycursor = mydb.cursor()
print("Database is connected")


def recommendation_letter(full_name, job_role, start_date, salary, Hr_name):
    # Create a new Document
    doc = Document()

    # Add a heading
    doc.add_heading('Letter of Recommendation', level=1)

    # Add the current date
    current_date = datetime.now().strftime('%Y-%m-%d')
    doc.add_paragraph(f'Date: {current_date}')

    # Add recipient details
    doc.add_paragraph(f'Dear {full_name},')
    doc.add_paragraph(f'Re: Recommendation for {full_name}')

    # Add body of the letter
    doc.add_paragraph(f'I am writing to highly recommend {full_name} for {specific_role}. I have had the pleasure of working closely with Her for {duration}, and during this time, I have been consistently impressed by his/her/their exceptional {skills}, dedication, and contributions.')

    doc.add_paragraph(f'{full_name} is an {qualities} individual who possesses a unique combination of {skills}. His/Her/Their ability to {accomplishments} demonstrates a keen intellect, strong work ethic, and a commitment to excellence.')

    doc.add_paragraph(f'One of {full_name}\'s most commendable qualities is his/her/their {quality} skills. Whether collaborating with colleagues, leading a team, or communicating with clients, he/she/they consistently exhibit {skills} communication skills that foster a positive and productive working environment.')

    doc.add_paragraph(f'Furthermore, {full_name} is highly {qualities}, enabling Her to tackle challenges with a creative and strategic mindset. Her ability to  showcases Her capacity to think critically and find effective solutions even in high-pressure situations.')

    doc.add_paragraph(f'In addition to his/her/their technical skills, {full_name} is a person of integrity and reliability. He/She/They approach his/her/their work with a high level of professionalism and always meet deadlines with a sense of responsibility. I have witnessed firsthand his/her/their commitment to {ethical_considerations}, making him/her/them a trustworthy and dependable team member.')

    # Add a closing statement
    doc.add_paragraph('I am confident that {full_name} will excel in {specific_role} and will bring a unique perspective and valuable contributions to your team. I wholeheartedly recommend {him/her/them} without reservation, and I believe {he/she/they} will be a valuable asset to {recipient_title}.')

    # Add a closing salutation
    doc.add_paragraph(f'Please feel free to contact me at [your email address] or [your phone number] if you require any further information. Thank you for considering my recommendation.')

    doc.add_paragraph(f'Sincerely,')
    doc.add_paragraph(f'[Your Full Name]')
   

    # Save the document
    doc_filename = f'Recommendation_{full_name.replace(" ","_")}.docx'
    doc.save(doc_filename)

    convert(doc_filename)

# Take user input for user id
mycursor.execute("SELECT * FROM data.internship")
results = mycursor.fetchall()

for result in results:
    recommendation_letter(str(result[0]), result[1], result[2], result[3], result[4])


mydb.close()