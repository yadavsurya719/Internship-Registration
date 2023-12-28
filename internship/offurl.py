from docx import Document
from datetime import datetime
from docx2pdf import convert
import pymysql
import os

# Connect to the MySQL database
mydb = pymysql.connect(host='localhost', user='root', password='surya7993290146', database='data')
mycursor = mydb.cursor()

print("Database is connected")

def offer_letter_to_pdf(full_name, job_role, start_date, salary, Hr_name):
    # Create a new Document
    doc = Document()

    # Add a heading
    doc.add_heading('Offer letter', level=1)

    # Add the current date
    current_date = datetime.now().strftime('%Y-%m-%d')
    doc.add_paragraph(f'Date: {current_date}')

    # Add employee details
    doc.add_paragraph(f'Dear {full_name},')
    doc.add_paragraph(f'We at {full_name} are delighted to inform you that you have been selected to our company as a {job_role}. Your joining date is {start_date} and you have to report to the office at 9:00 AM. Your salary will be {salary} per month.')

    # Calculate the number of days
    doc.add_paragraph('Regards,')

    # Add a closing statement
    doc.add_paragraph(f'{Hr_name}[HR], Pragmatiq Systems')

    # Generate a unique file name based on the user ID and name
    doc_filename = f'Offer_Letter_{full_name.replace(" ","_")}.docx'

    # Save the Word document
    doc.save(doc_filename)

    # Convert the document to PDF
    pdf_filename = f'Offer_Letter_{full_name.replace(" ","_")}.pdf'
    convert(doc_filename, pdf_filename)

    # Assuming the PDFs are stored in a directory called 'pdfs' on the server
    pdf_url = f'http://127.0.0.1:5000/pdfs/{pdf_filename}'

    return pdf_filename, offer_letter

# Iterate through the internship data and generate offer letters
mycursor.execute("SELECT * FROM data.internship")
results = mycursor.fetchall()

for result in results:
    user_id = result[0]
    full_name = result[1]
    job_role = result[2]
    start_date = result[3]
    salary = result[4]
    hr_name = result[5]

    # Generate the offer letter and get the PDF filename and URL
    pdf_filename, offer_letter = offer_letter_to_pdf(full_name, job_role, start_date, salary, hr_name)

    # Update the database with the PDF URL
    update_query = "UPDATE data.internship SET pdf_url = %s WHERE id = %s"
    mycursor.execute(update_query, (offer_letter, user_id))
    mydb.commit()

mydb.close()
